import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

base = r'd:\[PUB] India_runs_data_and_ai_challenge\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge'

for fname in ['job_description.docx', 'submission_spec.docx', 'redrob_signals_doc.docx', 'README.docx']:
    fpath = os.path.join(base, fname)
    print(f"\n{'='*80}")
    print(f"FILE: {fname}")
    print(f"{'='*80}")
    doc = Document(fpath)
    for p in doc.paragraphs:
        print(p.text)
    
    # Also extract tables if any
    for i, table in enumerate(doc.tables):
        print(f"\n--- TABLE {i+1} ---")
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            print(" | ".join(cells))
