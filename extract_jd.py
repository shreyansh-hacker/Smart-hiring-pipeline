import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'd:\[PUB] India_runs_data_and_ai_challenge\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge\job_description.docx')
for p in doc.paragraphs:
    print(p.text)
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        cells = [cell.text for cell in row.cells]
        print(" | ".join(cells))
