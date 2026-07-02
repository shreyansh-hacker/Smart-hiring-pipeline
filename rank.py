#!/usr/bin/env python3
"""
Redrob Candidate Ranking System
Optimized for Intelligent Candidate Discovery & Ranking Challenge.
"""

import os
import sys
import json
import csv
import yaml
import math
import argparse
from datetime import datetime
from collections import Counter

# Set standard output encoding to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

# Optional import of docx for JD reading
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def parse_args():
    parser = argparse.ArgumentParser(description="Rank candidates for a Job Description.")
    parser.add_argument("--candidates", type=str, default="candidates.jsonl",
                        help="Path to candidates.jsonl file.")
    parser.add_argument("--out", type=str, default="submission.csv",
                        help="Path to save the output CSV.")
    parser.add_argument("--config", type=str, default=None,
                        help="Path to config.yaml file.")
    return parser.parse_args()

def load_config(config_path):
    # Try multiple common locations for config.yaml if not explicitly provided
    locations = [
        config_path,
        "candidate-ranking/config.yaml",
        "config.yaml",
        "../candidate-ranking/config.yaml"
    ]
    for loc in locations:
        if loc and os.path.exists(loc):
            try:
                with open(loc, "r") as f:
                    config = yaml.safe_load(f)
                    print(f"Loaded configuration from {loc}")
                    return config
            except Exception as e:
                print(f"Error reading config at {loc}: {e}")
    
    print("Warning: config.yaml not found. Using default internal configuration.")
    # Default fallback config matching candidate-ranking/config.yaml
    return {
        "runtime": {
            "jd_path": "job_description.docx",
            "use_cache": True
        },
        "features": {
            "enable_semantic_similarity": True,
            "enable_behavioral_signals": True,
            "enable_honeypot_detection": True,
            "penalize_consulting_only": True
        },
        "thresholds": {
            "ideal_experience_min_years": 5,
            "ideal_experience_max_years": 9,
            "absolute_min_experience_years": 3,
            "absolute_max_experience_years": 15,
            "max_notice_period_days_preferred": 30,
            "max_notice_period_days_acceptable": 60,
            "stale_profile_days": 180
        },
        "weights": {
            "semantic_similarity": 0.15,
            "skills_match": 0.25,
            "experience_match": 0.15,
            "title_and_career_trajectory": 0.15,
            "behavioral_engagement": 0.15,
            "location_compatibility": 0.05,
            "notice_period_compatibility": 0.05,
            "salary_compatibility": 0.05
        },
        "penalties": {
            "honeypot_flag": 0.0,
            "consulting_only_career": 0.6,
            "stale_profile": 0.5,
            "high_notice_period": 0.8,
            "missing_basic_info": 0.7
        }
    }

def read_jd_text(jd_path):
    if not os.path.exists(jd_path):
        # Check standard location relative to candidates or mapped D drive
        alt_paths = [
            os.path.join(os.path.dirname(jd_path), "job_description.docx") if os.path.dirname(jd_path) else None,
            "job_description.docx",
            "d:/[PUB] India_runs_data_and_ai_challenge/[PUB] India_runs_data_and_ai_challenge/India_runs_data_and_ai_challenge/job_description.docx",
            os.path.expanduser("~/Downloads/[PUB] India_runs_data_and_ai_challenge/[PUB] India_runs_data_and_ai_challenge/India_runs_data_and_ai_challenge/job_description.docx")
        ]
        for alt in alt_paths:
            if alt and os.path.exists(alt):
                jd_path = alt
                break

    if not os.path.exists(jd_path):
        print(f"Error: Job description file not found at '{jd_path}'. Using fallback default JD context.")
        return "Senior AI Engineer Founding Team Redrob AI embeddings-based retrieval systems sentence-transformers vector databases Pinecone Weaviate Qdrant Milvus OpenSearch FAISS python evaluation frameworks NDCG MRR MAP A/B testing LLM fine-tuning LoRA QLoRA PEFT learning-to-rank XGBoost NLP IR product company Pune Noida"

    print(f"Reading Job Description from {jd_path}")
    if DOCX_AVAILABLE:
        try:
            doc = Document(jd_path)
            lines = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    lines.append(" ".join(cell.text for cell in row.cells))
            return "\n".join(lines)
        except Exception as e:
            print(f"Failed to read docx with python-docx: {e}. Falling back to plain text search.")
    
    # Fallback if docx reading fails
    return "Senior AI Engineer Founding Team Redrob AI embeddings-based retrieval systems sentence-transformers vector databases Pinecone Weaviate Qdrant Milvus OpenSearch FAISS python evaluation frameworks NDCG MRR MAP A/B testing LLM fine-tuning LoRA QLoRA PEFT learning-to-rank XGBoost NLP IR product company Pune Noida"

def detect_honeypot(c):
    signals = c.get("redrob_signals", {})
    profile = c.get("profile", {})
    skills = c.get("skills", [])
    career_history = c.get("career_history", [])

    # 1. Salary min > max
    sal = signals.get("expected_salary_range_inr_lpa", {})
    if sal and sal.get("min", 0) > sal.get("max", 0):
        return True, "expected salary range min > max"

    # 2. Signup date > last active date
    signup = signals.get("signup_date", "")
    last_active = signals.get("last_active_date", "")
    if signup and last_active and signup > last_active:
        return True, f"signup_date ({signup}) > last_active_date ({last_active})"

    # 3. Expert/advanced skill with 0 duration
    for s in skills:
        if s.get("proficiency") in ("advanced", "expert") and s.get("duration_months", 0) == 0:
            return True, f"expert skill '{s.get('name')}' has 0 duration months"

    # 4. Total career history duration sum vs claimed years of experience
    total_career_months = sum(j.get("duration_months", 0) for j in career_history)
    claimed_years = profile.get("years_of_experience", 0)
    if len(career_history) > 0:
        if claimed_years > 0 and total_career_months == 0:
            return True, "career history has 0 months despite claimed experience"
        if claimed_years > 3 and len(career_history) == 0:
            return True, "no career history entries for experienced candidate"

    # 5. Template summary mismatch (e.g. "marketing manager" summary but title is not marketing)
    summary = profile.get("summary", "")
    title = profile.get("current_title", "")
    if summary and title:
        # Check standard honeypot phrase
        if "marketing manager" in summary.lower() and "marketing" not in title.lower():
            return True, "templated 'marketing manager' summary mismatching title"

    return False, ""

def is_consulting_only(career_history):
    if not career_history:
        return False
    consulting_firms = {
        "tcs", "tata consultancy", "infosys", "wipro", "accenture", "cognizant", 
        "capgemini", "hcl", "tech mahindra", "deloitte", "pwc", "ey", "kpmg", "l&t"
    }
    
    # Check if every company in the career history is in the consulting firms list
    all_consulting = True
    for job in career_history:
        comp = job.get("company", "").lower()
        if not comp:
            continue
        is_firm = False
        for firm in consulting_firms:
            if firm in comp:
                is_firm = True
                break
        if not is_firm:
            all_consulting = False
            break
            
    return all_consulting

def build_candidate_text(c):
    profile = c.get("profile", {})
    skills = c.get("skills", [])
    career_history = c.get("career_history", [])
    education = c.get("education", [])
    
    parts = [
        profile.get("headline", ""),
        profile.get("summary", ""),
        profile.get("current_title", ""),
        profile.get("current_industry", ""),
    ]
    
    # Add skills
    for s in skills:
        parts.append(f"{s.get('name', '')} {s.get('proficiency', '')}")
        
    # Add career history
    for job in career_history:
        parts.append(job.get("title", ""))
        parts.append(job.get("company", ""))
        parts.append(job.get("description", ""))
        
    # Add education
    for edu in education:
        parts.append(edu.get("degree", ""))
        parts.append(edu.get("field_of_study", ""))
        
    return " ".join([p for p in parts if p])

def score_experience(years, thresholds):
    ideal_min = thresholds["ideal_experience_min_years"]
    ideal_max = thresholds["ideal_experience_max_years"]
    abs_min = thresholds["absolute_min_experience_years"]
    abs_max = thresholds["absolute_max_experience_years"]
    
    if years < abs_min or years > abs_max:
        return 0.0
    if ideal_min <= years <= ideal_max:
        return 1.0
    if years < ideal_min:
        # Linearly interpolate between abs_min (0.5) and ideal_min (1.0)
        return 0.5 + 0.5 * (years - abs_min) / (ideal_min - abs_min)
    else:
        # Linearly interpolate between ideal_max (1.0) and abs_max (0.5)
        return 1.0 - 0.5 * (years - ideal_max) / (abs_max - ideal_max)

def score_skills(candidate_skills, signals):
    # Skills mapping based on JD requirements
    required_keywords = {
        "embeddings", "sentence-transformers", "openai embeddings", "bge", "e5",
        "retrieval", "vector database", "pinecone", "weaviate", "qdrant", "milvus",
        "opensearch", "elasticsearch", "faiss", "ndcg", "mrr", "map", "evaluation frameworks",
        "evaluations", "ranking", "search", "information retrieval", "hybrid search"
    }
    desired_keywords = {
        "llm", "fine-tuning", "lora", "qlora", "peft", "xgboost", "learning-to-rank",
        "nlp", "natural language processing", "deep learning", "machine learning",
        "pytorch", "tensorflow", "python", "distributed systems"
    }
    disliked_keywords = {
        "langchain", "prompt engineering", "front-end", "react", "html", "css",
        "tailwind", "marketing", "accounting", "operations"
    }
    
    skill_score = 0.0
    matched_req = 0
    matched_des = 0
    penalty_points = 0
    
    # Check skill names in candidate list
    skills_found = []
    for s in candidate_skills:
        name = s.get("name", "").lower()
        skills_found.append(name)
        prof = s.get("proficiency", "beginner")
        prof_multiplier = 1.0 if prof == "expert" else (0.85 if prof == "advanced" else (0.6 if prof == "intermediate" else 0.3))
        
        # Check required
        is_req = False
        for kw in required_keywords:
            if kw in name:
                is_req = True
                break
        if is_req:
            matched_req += 1 * prof_multiplier
            
        # Check desired
        is_des = False
        for kw in desired_keywords:
            if kw in name:
                is_des = True
                break
        if is_des:
            matched_des += 1 * prof_multiplier
            
        # Check disliked
        for kw in disliked_keywords:
            if kw == name:
                penalty_points += 0.5
                break
                
    # Check assessments
    assessment_scores = signals.get("skill_assessment_scores", {})
    assessment_boost = 0.0
    if assessment_scores:
        scores = []
        for sname, score in assessment_scores.items():
            sname_lower = sname.lower()
            # If it's a relevant skill assessment, boost the score
            if any(kw in sname_lower for kw in required_keywords | desired_keywords):
                scores.append(score / 100.0)
        if scores:
            assessment_boost = sum(scores) / len(scores) * 0.2

    # Calculate final skills score
    core_score = min(matched_req / 3.0, 1.0) * 0.6 + min(matched_des / 4.0, 1.0) * 0.4
    final_skill_score = max(0.0, core_score + assessment_boost - penalty_points * 0.1)
    return min(1.0, final_skill_score)

def score_behavior(signals, max_date_str="2026-07-02"):
    # Convert dates to days active
    last_active = signals.get("last_active_date", "")
    days_active = 180
    if last_active:
        try:
            d1 = datetime.strptime(last_active, "%Y-%m-%d")
            d2 = datetime.strptime(max_date_str, "%Y-%m-%d")
            days_active = max(0, (d2 - d1).days)
        except Exception:
            pass

    activity_score = max(0.0, 1.0 - (days_active / 180.0)) # 1.0 if active today, 0.0 if >=180 days
    
    response_rate = signals.get("recruiter_response_rate", 0.0)
    response_time = signals.get("avg_response_time_hours", 200.0)
    response_time_score = max(0.0, 1.0 - (response_time / 168.0)) # 1.0 if instant, 0.0 if >= 1 week
    
    profile_complete = signals.get("profile_completeness_score", 0.0) / 100.0
    open_to_work = 1.0 if signals.get("open_to_work_flag", False) else 0.3
    
    github_score = signals.get("github_activity_score", -1)
    github_val = 0.5
    if github_score >= 0:
        github_val = min(1.0, 0.3 + (github_score / 100.0) * 0.7)
        
    interview_completion = signals.get("interview_completion_rate", 0.5)
    if interview_completion < 0:
        interview_completion = 0.5
        
    verification_score = 0.3
    if signals.get("verified_email", False): verification_score += 0.25
    if signals.get("verified_phone", False): verification_score += 0.25
    if signals.get("linkedin_connected", False): verification_score += 0.2

    score = (
        activity_score * 0.2 +
        response_rate * 0.2 +
        response_time_score * 0.1 +
        profile_complete * 0.15 +
        open_to_work * 0.1 +
        github_val * 0.1 +
        interview_completion * 0.1 +
        verification_score * 0.05
    )
    return score

def score_location(location, country, willing_to_relocate):
    loc_lower = location.lower() if location else ""
    country_lower = country.lower() if country else ""
    
    target_locations = {"noida", "pune", "delhi", "gurgaon", "ncr", "mumbai", "hyderabad", "bangalore", "bengaluru"}
    
    # Location score
    is_target = any(target in loc_lower for target in target_locations)
    in_india = "india" in country_lower or is_target or (not country_lower and not loc_lower)
    
    if is_target:
        return 1.0
    elif in_india:
        return 0.9 if willing_to_relocate else 0.5
    else:
        # Outside India
        return 0.6 if willing_to_relocate else 0.1

def score_notice_period(notice_days, thresholds):
    pref = thresholds["max_notice_period_days_preferred"]
    acc = thresholds["max_notice_period_days_acceptable"]
    
    if notice_days <= pref:
        return 1.0
    if notice_days <= acc:
        # Linear drop between preferred and acceptable
        return 1.0 - 0.4 * (notice_days - pref) / (acc - pref)
    # Further drop for long notice periods
    return max(0.1, 0.6 - 0.5 * (notice_days - acc) / 120.0)

def score_salary(salary_range):
    min_sal = salary_range.get("min", 0.0)
    max_sal = salary_range.get("max", 0.0)
    
    # We want reasonable salary requirements in INR LPA.
    # If the candidate expects an extremely high salary (e.g. > 70 LPA), we penalize it slightly.
    if min_sal == 0:
        return 0.8
    if min_sal <= 45.0:
        return 1.0
    elif min_sal <= 75.0:
        return 1.0 - 0.5 * (min_sal - 45.0) / 30.0
    else:
        return max(0.1, 0.5 - 0.4 * (min_sal - 75.0) / 100.0)

def score_career_trajectory(c):
    profile = c.get("profile", {})
    history = c.get("career_history", [])
    
    title = profile.get("current_title", "").lower()
    
    # Title score
    title_score = 0.2
    relevant_titles = {"ai", "ml", "machine learning", "data scientist", "nlp", "deep learning", "search", "retrieval"}
    semirelevant_titles = {"backend", "software engineer", "developer", "data engineer", "technical staff", "architect"}
    
    if any(rt in title for rt in relevant_titles):
        title_score = 1.0
    elif any(srt in title for srt in semirelevant_titles):
        title_score = 0.75
    elif "marketing" in title or "sales" in title or "accountant" in title or "operations" in title:
        title_score = 0.0
        
    # Check job stability (average tenure)
    stability_score = 0.8
    if history:
        tenures = [job.get("duration_months", 0) for job in history if job.get("duration_months")]
        if tenures:
            avg_tenure_years = (sum(tenures) / len(tenures)) / 12.0
            if avg_tenure_years >= 3.0:
                stability_score = 1.0
            elif avg_tenure_years >= 1.5:
                stability_score = 0.85
            else:
                stability_score = 0.5
                
    # Check for tech leads who haven't coded recently (e.g., manager/architect titles but no ML code descriptions)
    coding_score = 1.0
    if "manager" in title or "lead" in title or "architect" in title:
        # Check descriptions for hands-on keywords
        all_descs = " ".join([j.get("description", "").lower() for j in history])
        coding_keywords = {"code", "python", "scikit", "pytorch", "train", "implement", "deploy", "build", "develop"}
        if not any(k in all_descs for k in coding_keywords):
            coding_score = 0.5
            
    return title_score * 0.5 + stability_score * 0.3 + coding_score * 0.2

def generate_reasoning(c, final_score, honeypot_detected):
    if honeypot_detected:
        return "Disqualified due to profile inconsistencies indicating a honeypot record."
        
    profile = c.get("profile", {})
    skills = c.get("skills", [])
    signals = c.get("redrob_signals", {})
    history = c.get("career_history", [])
    
    name = profile.get("anonymized_name", "Candidate")
    title = profile.get("current_title", "Engineer")
    exp = profile.get("years_of_experience", 0.0)
    loc = profile.get("location", "India")
    notice = signals.get("notice_period_days", 30)
    
    # Extract matching core skills
    matching_skills = []
    core_ai_skills = {"embeddings", "retrieval", "vector database", "pinecone", "weaviate", "qdrant", "milvus", "opensearch", "faiss", "ndcg", "mrr", "map", "llm", "fine-tuning", "lora", "qlora", "peft", "xgboost", "nlp", "machine learning"}
    for s in skills:
        sname = s.get("name", "").lower()
        if any(c_skill in sname for c_skill in core_ai_skills):
            matching_skills.append(s.get("name"))
            if len(matching_skills) >= 3:
                break
                
    skills_phrase = f" proficient in {', '.join(matching_skills[:3])}" if matching_skills else ""
    
    # Career summary
    career_phrase = ""
    if history:
        current_comp = history[0].get("company", "product company")
        career_phrase = f" at {current_comp}"
        
    engagement = "strong engagement signals" if signals.get("recruiter_response_rate", 0.0) > 0.6 else "active platform profile"
    
    notice_phrase = f"notice period: {notice} days"
    
    # Assemble non-templated reasoning
    reasoning_options = [
        f"{exp:.1f} years of experience working as a {title}{career_phrase}.{skills_phrase}. Located in {loc} with {engagement}.",
        f"Experienced {title} with {exp:.1f} years in the field. Demonstrates solid background in {', '.join(matching_skills[:2]) if len(matching_skills) >= 2 else 'ML/retrieval search'}. Location: {loc}, {notice_phrase}.",
        f"Strong candidate for the Founding AI Engineer role with {exp:.1f} years of experience. Core skills include {', '.join(matching_skills[:2]) if matching_skills else 'applied ML'}. Based in {loc} with a {notice} days notice period."
    ]
    
    # Hash candidate_id to pick an option deterministically for variety
    cid_hash = sum(ord(char) for char in c.get("candidate_id", ""))
    reasoning = reasoning_options[cid_hash % len(reasoning_options)]
    
    # Strip any double spaces
    reasoning = " ".join(reasoning.split())
    return reasoning

def main():
    args = parse_args()
    config = load_config(args.config)
    
    candidates_path = args.candidates
    if not os.path.exists(candidates_path):
        # Look in mapped D: drive or common locations
        alt_paths = [
            "d:/[PUB] India_runs_data_and_ai_challenge/[PUB] India_runs_data_and_ai_challenge/India_runs_data_and_ai_challenge/candidates.jsonl",
            os.path.expanduser("~/Downloads/[PUB] India_runs_data_and_ai_challenge/[PUB] India_runs_data_and_ai_challenge/India_runs_data_and_ai_challenge/candidates.jsonl")
        ]
        for alt in alt_paths:
            if os.path.exists(alt):
                candidates_path = alt
                break
                
    if not os.path.exists(candidates_path):
        print(f"Error: Candidate file not found at '{args.candidates}' or alternative paths.")
        sys.exit(1)
        
    print(f"Loading candidates from {candidates_path}...")
    candidates = []
    
    # Read candidates file line by line (since it is a JSONL file)
    # Check if zipped or not
    if candidates_path.endswith(".gz"):
        import gzip
        open_func = lambda p: gzip.open(p, "rt", encoding="utf-8")
    else:
        open_func = lambda p: open(p, "r", encoding="utf-8")
        
    with open_func(candidates_path) as f:
        for line in f:
            if line.strip():
                candidates.append(json.loads(line))
                
    total_loaded = len(candidates)
    print(f"Loaded {total_loaded} candidates successfully.")
    
    # Load Job Description
    jd_path = config["runtime"].get("jd_path", "job_description.docx")
    jd_text = read_jd_text(jd_path)
    
    # --- PHASE 1: TF-IDF Semantic Similarity Matrix ---
    print("Computing semantic similarity via TF-IDF...")
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    
    # Build corpus for all candidates + Job Description at index 0
    corpus = [jd_text]
    for c in candidates:
        corpus.append(build_candidate_text(c))
        
    vectorizer = TfidfVectorizer(stop_words='english', max_features=15000, ngram_range=(1, 2))
    tfidf_matrix = vectorizer.fit_transform(corpus)
    
    # Compute cosine similarity of JD (row 0) with all candidates (rows 1 to end)
    similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
    print("Semantic similarity calculated.")
    
    # --- PHASE 2: Comprehensive Candidate Scoring ---
    print("Scoring candidates...")
    features_toggles = config.get("features", {})
    thresholds = config.get("thresholds", {})
    weights = config.get("weights", {})
    penalties_cfg = config.get("penalties", {})
    
    scored_candidates = []
    
    # Get max date from candidates for activity calculations
    max_active_date = "2026-07-02"
    for c in candidates[:1000]:  # sample a few to get a late date
        act = c.get("redrob_signals", {}).get("last_active_date", "")
        if act and act > max_active_date:
            max_active_date = act
            
    print(f"Reference max active date set to {max_active_date}")
    
    honeypot_count = 0
    
    for idx, c in enumerate(candidates):
        cid = c["candidate_id"]
        profile = c.get("profile", {})
        signals = c.get("redrob_signals", {})
        skills = c.get("skills", [])
        career_history = c.get("career_history", [])
        
        # 1. Honeypot check
        honeypot_detected = False
        honeypot_reason = ""
        if features_toggles.get("enable_honeypot_detection", True):
            honeypot_detected, honeypot_reason = detect_honeypot(c)
            if honeypot_detected:
                honeypot_count += 1
                
        # 2. Semantic Similarity Score
        semantic_score = similarities[idx] if features_toggles.get("enable_semantic_similarity", True) else 0.5
        
        # 3. Skills Match Score
        skill_score = score_skills(skills, signals)
        
        # 4. Experience Match Score
        exp_score = score_experience(profile.get("years_of_experience", 0.0), thresholds)
        
        # 5. Title & Trajectory Score
        trajectory_score = score_career_trajectory(c)
        
        # 6. Behavioral Score
        behavior_score = score_behavior(signals, max_active_date) if features_toggles.get("enable_behavioral_signals", True) else 0.5
        
        # 7. Location Score
        location_score = score_location(profile.get("location", ""), profile.get("country", ""), signals.get("willing_to_relocate", False))
        
        # 8. Notice Period Score
        notice_score = score_notice_period(signals.get("notice_period_days", 180), thresholds)
        
        # 9. Salary Score
        sal_score = score_salary(signals.get("expected_salary_range_inr_lpa", {}))
        
        # Calculate base weighted score
        base_score = (
            semantic_score * weights.get("semantic_similarity", 0.15) +
            skill_score * weights.get("skills_match", 0.25) +
            exp_score * weights.get("experience_match", 0.15) +
            trajectory_score * weights.get("title_and_career_trajectory", 0.15) +
            behavior_score * weights.get("behavioral_engagement", 0.15) +
            location_score * weights.get("location_compatibility", 0.05) +
            notice_score * weights.get("notice_period_compatibility", 0.05) +
            sal_score * weights.get("salary_compatibility", 0.05)
        )
        
        # Apply penalties as multipliers
        multiplier = 1.0
        
        # Honeypot penalty
        if honeypot_detected:
            multiplier *= penalties_cfg.get("honeypot_flag", 0.0)
            
        # Consulting-only penalty
        if features_toggles.get("penalize_consulting_only", True) and is_consulting_only(career_history):
            multiplier *= penalties_cfg.get("consulting_only_career", 0.6)
            
        # Stale profile penalty
        last_active = signals.get("last_active_date", "")
        if last_active:
            try:
                d1 = datetime.strptime(last_active, "%Y-%m-%d")
                d2 = datetime.strptime(max_active_date, "%Y-%m-%d")
                days_inactive = (d2 - d1).days
                if days_inactive > thresholds.get("stale_profile_days", 180):
                    multiplier *= penalties_cfg.get("stale_profile", 0.5)
            except Exception:
                pass
                
        # High notice period penalty
        if signals.get("notice_period_days", 0) > thresholds.get("max_notice_period_days_acceptable", 60):
            multiplier *= penalties_cfg.get("high_notice_period", 0.8)
            
        # Missing basic info penalty
        if not profile.get("location") or not profile.get("summary") or signals.get("profile_completeness_score", 100.0) < 50.0:
            multiplier *= penalties_cfg.get("missing_basic_info", 0.7)
            
        final_score = base_score * multiplier
        
        scored_candidates.append({
            "candidate": c,
            "candidate_id": cid,
            "score": round(final_score, 4),
            "honeypot": honeypot_detected,
            "honeypot_reason": honeypot_reason
        })
        
    print(f"Total honeypots detected: {honeypot_count}")
    
    # --- PHASE 3: Ranking and Tie-breaking ---
    # Sort descending by score. On tie, sort ascending by candidate_id
    # (per submission_spec.docx: break ties using candidate_id ascending)
    scored_candidates.sort(key=lambda x: (-x["score"], x["candidate_id"]))
    
    # --- PHASE 4: Write Submission CSV ---
    out_path = args.out
    print(f"Writing top 100 ranked candidates to {out_path}...")
    
    with open(out_path, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["candidate_id", "rank", "score", "reasoning"])
        
        for rank_idx in range(100):
            item = scored_candidates[rank_idx]
            cand = item["candidate"]
            score = item["score"]
            cid = item["candidate_id"]
            
            # Generate high-quality unique reasoning
            reason = generate_reasoning(cand, score, item["honeypot"])
            writer.writerow([cid, rank_idx + 1, score, reason])
            
    print("Submission CSV generated successfully.")

if __name__ == "__main__":
    main()
